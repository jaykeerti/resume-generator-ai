"""
Document Parser Service
Handles extraction of text from PDF, DOCX, and TXT files
"""

import io
import os
from typing import BinaryIO
import logging

# PDF parsing
try:
    from pypdf import PdfReader
except ImportError:
    from PyPDF2 import PdfReader

# DOCX parsing
from docx import Document

logger = logging.getLogger(__name__)


class DocumentParser:
    """Service for parsing various document formats"""

    def parse(self, content: bytes, filename: str) -> str:
        """
        Parse document content based on file extension

        Args:
            content: Raw file bytes
            filename: Original filename with extension

        Returns:
            Extracted text content

        Raises:
            ValueError: If file format is unsupported
            Exception: If parsing fails
        """
        file_ext = os.path.splitext(filename)[1].lower()

        try:
            if file_ext == ".pdf":
                return self._parse_pdf(content)
            elif file_ext == ".docx":
                return self._parse_docx(content)
            elif file_ext == ".txt":
                return self._parse_txt(content)
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")

        except Exception as e:
            logger.error(f"Error parsing {filename}: {str(e)}")
            raise

    def _parse_pdf(self, content: bytes) -> str:
        """
        Extract text from PDF file

        Args:
            content: PDF file bytes

        Returns:
            Extracted text
        """
        try:
            pdf_file = io.BytesIO(content)
            reader = PdfReader(pdf_file)

            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            extracted_text = "\n\n".join(text_parts)

            if not extracted_text.strip():
                raise ValueError("No text could be extracted from PDF. The file may be image-based or corrupted.")

            return extracted_text.strip()

        except Exception as e:
            logger.error(f"PDF parsing error: {str(e)}")
            raise ValueError(f"Failed to parse PDF: {str(e)}")

    def _parse_docx(self, content: bytes) -> str:
        """
        Extract text from DOCX file

        Args:
            content: DOCX file bytes

        Returns:
            Extracted text
        """
        try:
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)

            text_parts = []

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            extracted_text = "\n".join(text_parts)

            if not extracted_text.strip():
                raise ValueError("No text could be extracted from DOCX file.")

            return extracted_text.strip()

        except Exception as e:
            logger.error(f"DOCX parsing error: {str(e)}")
            raise ValueError(f"Failed to parse DOCX: {str(e)}")

    def _parse_txt(self, content: bytes) -> str:
        """
        Extract text from TXT file

        Args:
            content: TXT file bytes

        Returns:
            Decoded text
        """
        try:
            # Try UTF-8 first
            try:
                text = content.decode("utf-8")
            except UnicodeDecodeError:
                # Fallback to latin-1
                text = content.decode("latin-1")

            if not text.strip():
                raise ValueError("Text file is empty.")

            return text.strip()

        except Exception as e:
            logger.error(f"TXT parsing error: {str(e)}")
            raise ValueError(f"Failed to parse text file: {str(e)}")
